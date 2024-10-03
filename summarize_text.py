import os
import subprocess
import streamlit as st
from docx import Document
from fpdf import FPDF


def summarize_text(text, lang='en'):
    # Define the prompt
    prompt = f"""
    The following text is in its original language. Provide the output in this language: {lang}. 
    Format the output as follows:

    Summary:
    short summary of the video

    Key Takeaways:
    succinct bullet point list of key takeaways

    input text: {text}
    """

    # Call the Ollama model using subprocess
    command = ["ollama", "run", "llama3.1:8b", prompt]
    
    # Run the command and capture stdout and stderr
    result = subprocess.run(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding='utf-8'
    )

    # Collect the output, filtering out unwanted lines
    output_lines = result.stdout.splitlines()  # Split stdout into lines
    filtered_output = "\n".join(line for line in output_lines if "failed to get console mode" not in line)  # Filter out unwanted lines

    # Check for errors in stderr
    if result.returncode != 0:
        return f"Failed to generate summary. Error: {result.stderr.strip()}"

    # Return the cleaned output
    return filtered_output.strip()


def add_summary_to_document(video_title, summary_text, filename='video_summaries.docx'):
    # Create the document in the current directory
    doc = Document()
    doc.add_heading('Video Summaries', level=1)

    # Add the video title as a heading
    doc.add_heading(video_title, level=2)

    # Split the summary into points
    points = summary_text.split('\n')
    for point in points:
        doc.add_paragraph(point, style='ListBullet')

    # Save the document in the current directory
    doc.save(filename)
    return filename


def create_pdf_from_docx(docx_filename, pdf_filename):
    # Create a PDF from the DOCX file
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Read the DOCX file and add content to PDF
    doc = Document(docx_filename)
    for para in doc.paragraphs:
        pdf.multi_cell(0, 10, para.text)

    pdf.output(pdf_filename)


def main():
    st.title("YouTube Video Summarizer")

    video_title = st.text_input("Enter the video title:")
    text_to_summarize = st.text_area("Enter the text to summarize:")
    lang = st.text_input("Enter the language for the summary:")

    if st.button("Generate Summary"):
        summary = summarize_text(text_to_summarize, lang)
        st.write("Summary:")
        st.write(summary)

        # Add summary to the document and get the filename
        docx_filename = add_summary_to_document(video_title, summary)

        # Create a PDF version of the document
        pdf_filename = docx_filename.replace('.docx', '.pdf')
        create_pdf_from_docx(docx_filename, pdf_filename)

        # Provide a clickable link for DOCX
        st.markdown(f"[click here to download the video summaries]({docx_filename})", unsafe_allow_html=True)

        # Provide a download button for the DOCX document
        with open(docx_filename, "rb") as f:
            st.download_button(
                label="Download Summary DOCX",
                data=f,
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # Provide a download button for the PDF document
        with open(pdf_filename, "rb") as f:
            st.download_button(
                label="Download Summary PDF",
                data=f,
                file_name=pdf_filename,
                mime="application/pdf"
            )


if __name__ == "__main__":
    main()
